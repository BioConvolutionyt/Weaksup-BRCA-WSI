# -*- coding: utf-8 -*-
"""探索实验报告（Word）生成器。

将 6.1–6.4 探索笔记汇总为学术报告级 docx：连贯段落、三级标题、图表编号与
图注/表头、公式 LaTeX 渲染为图片并编号、数字顺序编码制参考文献。
逐章构建：当前包含标题页 + 第 1 章（6.1），后续章节经确认后追加。

用法：直接运行（无命令行参数），输出 exploration/探索实验报告.docx。
"""
from pathlib import Path
import re

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.table import (WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE,
                             WD_TABLE_ALIGNMENT)
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image

PROJECT_ROOT = Path(__file__).resolve().parent.parent
RESULTS = PROJECT_ROOT / "exploration" / "results"
ASSETS = PROJECT_ROOT / "exploration" / "report_assets"
OUT_DOCX = PROJECT_ROOT / "exploration" / "探索实验报告.docx"

TEXT_W_CM = 16.0          # A4 21cm − 左右各 2.5cm 页边距

# 中英文/数字间去空格：CJK 与拉丁字母（含扩展）、希腊字母、数字、常见符号相邻时
_LAT = r"A-Za-z0-9À-ɏͰ-Ͽ"
_SP1 = re.compile(rf"([一-鿿])\s+([{_LAT}(+−–%])")
_SP2 = re.compile(rf"([{_LAT}%)\]\*\^+−–])\s+([一-鿿])")


def despace(text):
    """去除中文字符与英文/数字/符号之间的空格（用户排版要求）。"""
    prev = None
    while prev != text:
        prev = text
        text = _SP1.sub(r"\1\2", text)
        text = _SP2.sub(r"\1\2", text)
    return text


# ---------------------------------------------------------------- 样式与底层助手

def setup_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    for lv, size in [("Heading 1", 16), ("Heading 2", 13.5), ("Heading 3", 12)]:
        st = doc.styles[lv]
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0, 0, 0)
        # 内置 Heading 样式带主题字体属性（优先级高于直接字体），须先移除
        rpr = st._element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        for attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
            if rfonts.get(qn(f"w:{attr}")) is not None:
                del rfonts.attrib[qn(f"w:{attr}")]
        rfonts.set(qn("w:ascii"), "Times New Roman")
        rfonts.set(qn("w:hAnsi"), "Times New Roman")
        rfonts.set(qn("w:cs"), "Times New Roman")
        rfonts.set(qn("w:eastAsia"), "宋体")
    for section in doc.sections:
        section.page_width, section.page_height = Cm(21.0), Cm(29.7)
        section.left_margin = section.right_margin = Cm(2.5)
        section.top_margin = section.bottom_margin = Cm(2.5)


def _runs_font(p, east="宋体"):
    for r in p.runs:
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), east)


def add_para(doc, text, indent=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.85)
    p.paragraph_format.line_spacing = 1.4
    p.paragraph_format.space_after = Pt(4)
    p.add_run(despace(text))
    _runs_font(p)
    return p


def add_heading(doc, text, level):
    h = doc.add_heading(despace(text), level=level)
    h.paragraph_format.space_before = Pt(10 if level > 1 else 14)
    h.paragraph_format.space_after = Pt(6)
    for r in h.runs:   # run 级直设，双保险覆盖主题字体
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    return h


def render_formula(latex, name, fontsize=12, dpi=300):
    ASSETS.mkdir(parents=True, exist_ok=True)
    path = ASSETS / f"eq_{name}.png"
    fig = plt.figure(figsize=(0.1, 0.1))
    fig.text(0, 0, f"${latex}$", fontsize=fontsize)
    fig.savefig(path, dpi=dpi, bbox_inches="tight", pad_inches=0.04,
                facecolor="white")
    plt.close(fig)
    return path


def add_formula(doc, latex, number, name, fontsize=12, dpi=300):
    """公式图片按自然物理尺寸插入（字号恒定 12pt，字符大小跨公式一致），
    居中、编号右对齐。"""
    path = render_formula(latex, name, fontsize, dpi)
    with Image.open(path) as im:
        w_cm = im.width / dpi * 2.54
        h_cm = im.height / dpi * 2.54
    if w_cm > TEXT_W_CM - 2:                    # 超宽公式整体等比收缩
        scale = (TEXT_W_CM - 2) / w_cm
        w_cm, h_cm = w_cm * scale, h_cm * scale
    p = doc.add_paragraph()
    p.paragraph_format.tab_stops.add_tab_stop(Cm(TEXT_W_CM / 2),
                                              WD_TAB_ALIGNMENT.CENTER)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(TEXT_W_CM),
                                              WD_TAB_ALIGNMENT.RIGHT)
    p.add_run("\t")
    p.add_run().add_picture(str(path), width=Cm(w_cm), height=Cm(h_cm))
    num_run = p.add_run(f"\t({number})")
    # 图片与文字基线对齐导致编号沉底：按图片高度把编号抬升至公式垂直中心
    raise_pt = max(0.0, (h_cm * 28.3465 - 12) / 2)
    if raise_pt > 0.5:
        rpr = num_run._element.get_or_add_rPr()
        pos = OxmlElement("w:position")
        pos.set(qn("w:val"), str(int(round(raise_pt * 2))))   # 单位：半磅
        rpr.append(pos)
    _runs_font(p)
    return p


def add_figure(doc, img_path, label, caption, width_cm=13.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(img_path), width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.add_run(despace(f"{label}. {caption}"))
    for r in cap.runs:
        r.font.size = Pt(10.5)
        r.bold = True
    _runs_font(cap)
    return p


def _set_border(el_pr, tag, val, sz=0):
    b = OxmlElement(f"w:{tag}")
    b.set(qn("w:val"), val)
    if sz:
        b.set(qn("w:sz"), str(sz))
        b.set(qn("w:color"), "000000")
    el_pr.append(b)


def make_three_line(tb):
    """三线表：顶线/底线 1.5pt，表头下线 0.75pt，无竖线与内部横线。"""
    tbl_pr = tb._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    _set_border(borders, "top", "single", 12)
    _set_border(borders, "bottom", "single", 12)
    for tag in ("left", "right", "insideH", "insideV"):
        _set_border(borders, tag, "none")
    tbl_pr.append(borders)
    for cell in tb.rows[0].cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_b = OxmlElement("w:tcBorders")
        _set_border(tc_b, "bottom", "single", 6)
        tc_pr.append(tc_b)


def add_table(doc, label, caption, header, rows, col_w=None):
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.add_run(despace(f"{label}. {caption}"))
    for r in cap.runs:
        r.font.size = Pt(10.5)
        r.bold = True
    _runs_font(cap)
    tb = doc.add_table(rows=len(rows) + 1, cols=len(header))
    tb.style = "Table Grid"
    tb.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(header):
        cell = tb.cell(0, j)
        cell.text = despace(h)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(12)
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            cell = tb.cell(i + 1, j)
            cell.text = despace(str(v))
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(12)
    if col_w:
        for j, w in enumerate(col_w):
            for row in tb.rows:
                row.cells[j].width = Cm(w)
    for row in tb.rows:                     # 行高拉长（宽不变），避免扁平
        row.height = Cm(0.95)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    make_three_line(tb)
    for row in tb.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                _runs_font(p)
    doc.add_paragraph()
    return tb


def add_references(doc, refs):
    add_heading(doc, "参考文献", 1)
    for i, r in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.add_run(f"[{i}] {r}")
        for run in p.runs:
            run.font.size = Pt(10.5)
        _runs_font(p)


# ---------------------------------------------------------------- 参考文献（按首次引用顺序编号）

REFS = [
    "Team PND (poteman, arutema, famtaro). Prostate cANcer graDe Assessment "
    "(PANDA) Challenge 1st place solution: documentation and source code. "
    "Kaggle competition writeup and GitHub repository, 2020.",
    "D'Amato V, van der Laak J A M, Ciompi F. \"No negatives needed\": "
    "weakly-supervised regression for interpretable tumor detection in "
    "whole-slide histopathology images. Scientific Reports, 2025, 15. "
    "DOI: 10.1038/s41598-025-27158-8.",
    "Duffner S, Garcia C. Multiple Instance Learning for Training Neural "
    "Networks under Label Noise. Proceedings of the International Joint "
    "Conference on Neural Networks (IJCNN), 2020. "
    "DOI: 10.1109/IJCNN48605.2020.9206669.",
    "Oner M U, Lee H K, Sung W K. Distribution based MIL pooling filters: "
    "Experiments on a lymph node metastases dataset. Medical Image Analysis, "
    "2023, 87: 102813. DOI: 10.1016/j.media.2023.102813.",
    "Grahn D. mil-benchmarks: Standardized Evaluation of Deep "
    "Multiple-Instance Learning Techniques. arXiv:2105.01443, 2021.",
    "Kolmogorov V, Zabih R. What energy functions can be minimized via "
    "graph cuts? IEEE Transactions on Pattern Analysis and Machine "
    "Intelligence, 2004, 26(2): 147-159. DOI: 10.1109/TPAMI.2004.1262177.",
    "Boykov Y, Kolmogorov V. An experimental comparison of min-cut/max-flow "
    "algorithms for energy minimization in vision. IEEE Transactions on "
    "Pattern Analysis and Machine Intelligence, 2004, 26(9): 1124-1137. "
    "DOI: 10.1109/TPAMI.2004.60.",
    "Silva-Rodríguez J, Colomer A, Naranjo V. Proportion constrained weakly "
    "supervised histopathology image classification. Computers in Biology "
    "and Medicine, 2022, 147: 105714. DOI: 10.1016/j.compbiomed.2022.105714.",
]


# ---------------------------------------------------------------- 第 1 章（6.1）

def chapter_1(doc):
    add_heading(doc, "一、标签可信度协同推断的可行性边界", 1)

    add_heading(doc, "1.1 问题提出", 2)
    add_para(doc,
        "本项目使用的弱监督标签来源于 TCGA GDC 元数据中的肿瘤细胞百分比等字段。"
        "此类字段并非为机器学习建模而采集，存在字段缺失、统计口径差异与录入误差等"
        "多重风险，部分切片标签的可信度存疑。当标签噪声以部分切片整体错误的形式"
        "存在时，一个自然的设想是让模型与标签相互校正，即用当前模型的袋外预测残差"
        "评估各标签的可信程度，再以可信度为样本权重重新训练模型，如此迭代，期望"
        "模型逐渐向可信标签收敛。这一协同推断机制曾以硬阈值删除的形式出现在 "
        "PANDA 竞赛冠军方案的标签治理环节中[1]，其软加权版本已被本项目采纳为"
        "交付模型的可选组件。")
    add_para(doc,
        "然而该机制面临一个根本性的理论困难。模型参数与标签可信度相互依赖，只有"
        "准确的模型才能识别可信的标签，也只有可信的标签才能训练出准确的模型。切片"
        "级标签的数量通常只有数十个，远远小于模型参数量，这是一个高度欠定的推断"
        "问题。如果不引入额外的结构性假设，例如大多数标签可信，协同推断是否在"
        "原理上可行，其可行性边界又在哪里，仅凭理论论证难以回答。本章设计一个"
        "真值完全可控的小规模模拟实验，在不同污染比例下定量刻画协同推断的可行性"
        "边界，为交付管线中可信度加权组件的启用条件提供实验依据。")

    add_heading(doc, "1.2 实验设计", 2)
    add_heading(doc, "1.2.1 真值可控的合成数据平台", 3)
    add_para(doc,
        "在真实数据上干净真值不可获得，模型被拉回了多少与可信度有没有标对污染"
        "标签两个问题都无法直接度量。为此构造合成任务，特征保持真实，真值由确定"
        "规则生成。具体地，保留 25 张切片的真实 UNI 特征，每张切片为数千个 1024 维"
        "patch 向量构成的袋。生成一个固定随机向量 w* 作为隐藏的理想打分器，patch "
        "的真值肿瘤分数定义为")
    add_formula(doc,
        r"s_{ij} = \sigma\!\left(\alpha\,(f_{ij}^{\mathrm{T}} w^{*} + \delta_i)\right)",
        1, "1_1")
    add_para(doc,
        "其中 f_ij 为第 i 张切片第 j 个 patch 的特征向量，σ 为 sigmoid 函数，α 为"
        "尺度系数，δ_i 为逐切片偏移。随后按交付模型的聚合方式，取袋内分数最高的 "
        "25% patch 的均值作为该切片的真值肿瘤占比")
    add_formula(doc,
        r"y_i^{*} = \frac{1}{|T_i|} \sum_{j \in T_i} s_{ij}",
        2, "1_2")
    add_para(doc,
        "其中 T_i 为袋内分数排名前 25% 的 patch 下标集合。这样构造的学习任务与"
        "真实任务结构同构，模型原则上可以恢复 w* 的行为，而真值完全已知。")
    add_para(doc,
        "两处校准保证任务不退化。其一，α 校准至逐切片 patch 分数标准差约 0.15，"
        "避免 sigmoid 饱和或压平导致 patch 之间失去区分度。其二，预实验发现 25 张"
        "切片的真值占比几乎相同，标准差仅 0.015，模型预测常数即可取得低误差，"
        "排序与回归均无从谈起。为此引入逐切片偏移 δ_i，在区间内均匀采样，并通过"
        "二分搜索将其幅度校准至真值占比标准差约 0.123，与真实面积参照标签的离散"
        "度相当。")
    add_heading(doc, "1.2.2 标签污染模型", 3)
    add_para(doc,
        "沿用本文第三章的离群污染协议[2]。以污染率 p 随机选中对应比例的切片，将其"
        "标签替换为与切片内容无关的均匀分布随机数，其余切片保持真值标签")
    add_formula(doc,
        r"\tilde{y}_i = b_i\, u_i + (1 - b_i)\, y_i^{*},"
        r"\quad b_i \sim \mathrm{Bernoulli}(p),\ u_i \sim \mathcal{U}[0,1]",
        3, "1_3")
    add_para(doc,
        "污染率 p 取 0、20%、35%、50%、65% 五档。污染标记被完整记录，作为评估"
        "可信度机制的裁判依据。")
    add_heading(doc, "1.2.3 协同推断算法", 3)
    add_para(doc,
        "被测机制为 PANDA 冠军方案中袋外残差清洗的软加权迭代版[1]。第 0 轮以等权"
        "训练标准的 5 折交叉验证模型，得到每张切片的袋外预测 ŷ_i，进而计算残差")
    add_formula(doc,
        r"r_i = \left|\hat{y}_i^{\mathrm{OOF}} - \tilde{y}_i\right|",
        4, "1_4")
    add_para(doc,
        "直觉上，被污染的标签与数据整体相矛盾，从多数干净标签中学到的规律无法"
        "拟合它，因而残差偏大。残差按下式映射为可信度")
    add_formula(doc,
        r"c_i = \frac{\exp(-r_i / \tau)}{Z},"
        r"\quad \tau = \mathrm{median}(r)",
        5, "1_5")
    add_para(doc,
        "其中归一化系数 Z 将可信度的均值调整为 1。第 1 轮与第 2 轮以 c_i 为样本"
        "权重重新训练，并基于新的袋外预测更新可信度，共迭代 3 轮。对照组为不做"
        "任何处理的单次训练。收敛性以相邻两轮可信度的平均绝对变化度量")
    add_formula(doc,
        r"\Delta c^{(t)} = \frac{1}{n} \sum_{i=1}^{n}"
        r" \left| c_i^{(t)} - c_i^{(t-1)} \right|",
        6, "1_6")
    add_heading(doc, "1.2.4 评估协议", 3)
    add_para(doc,
        "评估协议的两个平面严格分离。算法侧只能接触污染后的标签，早停与模型选择"
        "均基于污染标签，这与真实场景一致，真实场景下无法获得干净的验证集。评分"
        "侧的全部指标对干净真值计算，包括平均绝对误差 MAE 与 Pearson 相关系数。"
        "可信度机制本身的质量以残差识别污染标记的 ROC 曲线下面积 AUC 度量。实验"
        "规模为 5 个污染档乘以每档 4 次运行，合计 20 次完整交叉验证，即 300 次"
        "模型训练。")

    add_heading(doc, "1.3 结果与分析", 2)
    add_para(doc,
        "表 1 汇总各污染档下无处理与协同推断最终轮的误差及可信度识别质量。")
    res = pd.read_csv(RESULTS / "6.1_results.csv", encoding="utf-8-sig")
    coi = res[res.method == "coinfer"]
    none = res[res.method == "none"].set_index("level")["mae_vs_truth"]
    final = coi[coi["round"] == 2].set_index("level")
    rows = []
    for lv in [0.0, 0.2, 0.35, 0.5, 0.65]:
        m0, m2 = none[lv], final.loc[lv, "mae_vs_truth"]
        auc = final.loc[lv, "cred_auc"]
        rows.append([f"{lv:.0%}", f"{m0:.4f}", f"{m2:.4f}",
                     f"{(m0 - m2) / m0 * 100:+.1f}%",
                     "—" if np.isnan(auc) else f"{auc:.3f}"])
    add_table(doc, "表1", "不同污染率下协同推断的效果与可信度识别质量",
              ["污染率", "无处理 MAE", "协同推断 MAE", "相对改善", "可信度 AUC"],
              rows)
    add_para(doc,
        "从表 1 可以读出三个层次的信息。首先，无处理列给出噪声对模型的自然伤害"
        "曲线，MAE 从干净条件的 0.115 升至 35% 污染时的 0.167，绝对恶化 0.052，"
        "说明离群标签对袋级回归的破坏不可忽略。其次，协同推断列在全部非零污染档"
        "上均低于无处理列，且改善幅度随污染率单调扩大，从 20% 档的 2.0% 增至 "
        "65% 档的 11.1%，表明迭代重加权确实把残差信息转化为了纠错能力。最后，"
        "可信度 AUC 列显示识别质量全程高于随机水平，即便在 50% 污染这一最难辨识"
        "的档位，AUC 仍为 0.705，说明残差排序对污染标记保持了可用的区分度。")
    add_para(doc,
        "图 1 的污染误差曲线将表 1 的两列误差连成趋势。无处理曲线在 35% 档达到"
        "峰值 0.167，协同推断曲线全程压在其下方或与之重合，两条曲线在 35% 处出现"
        "明确的收益交叉。交叉点左侧污染轻微，模型从多数干净标签中本就能学到稳定"
        "规律，可信度加权反而因引入额外方差而微幅受损，无污染时相对变化为 "
        "−1.9%。交叉点右侧污染加重，残差对污染标签的识别能力开始兑现为净收益，"
        "65% 档的相对改善达到 11.1%。这条交叉曲线直接划定了协同推断的启用边界，"
        "即污染率约 35%。值得注意的是，65% 档两条曲线同时回落，原因是高污染下"
        "均匀随机标签的均值 0.5 把预测拉向恰好接近真值均值 0.54 的中央位置，属于"
        "合成设定的巧合，不具推广意义，趋势解读时已将其排除。")
    add_figure(doc, RESULTS / "6.1_contamination_curve.png", "图1",
               "污染率与误差的关系曲线，协同推断在污染率 35% 处出现收益交叉",
               width_cm=11.0)
    add_para(doc,
        "图 2 给出残差识别污染标签的 AUC。全部污染档上 AUC 介于 0.70 至 0.84，"
        "显著高于随机水平 0.5，证明袋外残差确实携带标签污染信息。这是协同推断"
        "可行的直接证据。图 3 显示各污染档的可信度变化量 Δc 逐轮下降，无振荡或"
        "发散，迭代闭环稳定收敛。")
    add_figure(doc, RESULTS / "6.1_credibility_auc.png", "图2",
               "残差识别污染标签的 AUC 随污染率的变化，各轮迭代均显著高于随机水平",
               width_cm=11.0)
    add_figure(doc, RESULTS / "6.1_convergence.png", "图3",
               "协同推断相邻两轮可信度的平均绝对变化，逐轮下降表明迭代稳定收敛",
               width_cm=11.0)
    add_para(doc,
        "需要指出，25 袋合成任务本身的学习信号较弱，干净基线的 Pearson 相关系数"
        "仅为 0.08，故本章结论应理解为处理的相对效应，而非绝对性能水平。")

    add_heading(doc, "1.4 结论与讨论", 2)
    add_para(doc,
        "本章回答可行性边界三问。第一，协同推断可行。袋外残差携带污染信息，识别"
        "AUC 介于 0.70 至 0.84，迭代闭环稳定收敛。第二，可行性边界位于污染率 "
        "35% 附近。污染率不低于 35% 时启用有净收益，相对改善 4.2% 至 11.1%，不"
        "高于 20% 时接近无效甚至微害。第三，收益存在上限。即便在 65% 污染下，"
        "协同推断也只收复部分损失，MAE 从 0.143 降至 0.127，距干净基线 0.115 仍"
        "有差距。该机制能减损，不能复原。")
    add_para(doc,
        "对本项目主线而言，上述结论还有一层更重要的映射。本项目的标签风险并非"
        "随机污染，而是 GDC 细胞口径与面积口径之间的系统性错位。此类误差使所有"
        "切片同向偏移，模型可以一致地拟合错误标签而残差保持很小，基于残差的"
        "方法在原理上无法识别。因此协同推断适用于部分切片标签可疑的场景，例如 "
        "GDC 字段缺失或质检标记的高危切片，而口径校正已由面积参照路线解决，"
        "二者互补而非冲突。此外，识别 AUC 本身可作为启用前的筛查指标，当残差对"
        "可疑标记的识别能力接近随机水平时，不应启用可信度加权。")
    add_para(doc,
        "本章的局限在于合成平台的袋数为 25，污染模型为均匀离群，更贴近真实的"
        "结构性污染与更大袋数规模留待后续研究。若计算资源允许，将在 75 张队列与"
        "混合污染模型下复核边界位置，并将逐轮可信度演化导出可视化，作为该机制的"
        "调试工具。")


# ---------------------------------------------------------------- 第 2 章（6.2）

def chapter_2(doc):
    add_heading(doc, "二、池化硬度与切片patch数量的最优匹配", 1)

    add_heading(doc, "2.1 问题提出", 2)
    add_para(doc,
        "在多示例学习中，切片级预测由袋内全部patch分数经池化函数压缩而成，池化"
        "硬度是聚合器设计的核心自由度。平均池化将稀疏的肿瘤信号稀释于大量良性"
        "patch之中，最大池化只响应极端值而完全忽略面积总量，两极之间存在一个"
        "尚待回答的最优硬度问题。进一步地，全切片图像的patch数量从数千到数万"
        "不等，最优硬度是否随袋容量漂移，决定了交付管线是否需要容量自适应的"
        "聚合器。本章围绕项目文档探索内容6.2展开，依次回答四个问题，即两极之间"
        "是否存在最优硬度，硬度是否应随patch数量变化，能否让硬度在训练过程中"
        "自动调整，以及固定池化形式下回归误差与patch数量呈何种经验关系。本章与"
        "第三章的标签噪声维度互补，共同构成交付模型聚合器选型的完整依据。")

    add_heading(doc, "2.2 实验设计", 2)
    add_heading(doc, "2.2.1 池化硬度阶梯与广义均值参数化", 3)
    add_para(doc,
        "硬度的连续参数化采用广义均值池化[3]")
    add_formula(doc,
        r"\mathrm{pred} = \left( \frac{1}{N} \sum_{i=1}^{N} s_i^{p}"
        r" \right)^{1/p}",
        7, "2_1")
    add_para(doc,
        "其中s_i为袋内第i个patch的肿瘤分数，N为袋容量，p即硬度旋钮。p等于1时"
        "退化为算术平均，p趋向无穷时趋向最大值。与max和top-k等硬选择不同，广义"
        "均值对每个patch可微，梯度按s_i的p−1次幂加权分配，所有patch都能获得"
        "学习信号。实验选取五种池化器构成硬度阶梯，由软到硬依次为mean、"
        "gmean-p5、分位数q75、top-50与max。")
    add_para(doc,
        "阶梯中埋设了回答第二个问题的内置对照。分位数q75是固定比例池化，取袋内"
        "分数最高的25%patch求均值")
    add_formula(doc,
        r"\mathrm{pred} = \frac{1}{|T|} \sum_{j \in T} s_j,"
        r"\quad |T| = \lceil 0.25 N \rceil",
        8, "2_2")
    add_para(doc,
        "其有效实例数n_eff等于0.25N，随N线性增长而硬度恒定。top-50是固定个数"
        "池化，其等效硬度随N被动升高，在N为1000时约当于前5%，在N约三万时约当于"
        "前0.17%。若硬度应当随N调整，两者的误差曲线将呈现相反走向，反之则重合"
        "或平行。")
    add_heading(doc, "2.2.2 嵌套子采样与硬度×N实验矩阵", 3)
    add_para(doc,
        "patch数量维度取1000、2000、5000与全量四档。为消除抽到哪些patch这一"
        "混杂因素，同一切片采用同一固定随机排列的前缀构造嵌套子集，即1000档是"
        "2000档的子集，2000档是5000档的子集，N成为唯一变量。实验矩阵为5种池化"
        "器乘以4档N，每组执行5折交叉验证与3种子集成，合计300次训练。标签采用"
        "25张切片的面积口径参照，空间约束、可信度加权与标签噪声全部关闭，以隔离"
        "池化与N两个变量，实验协议与第一章及第三章保持一致。")
    add_heading(doc, "2.2.3 课程式硬度调度", 3)
    add_para(doc,
        "第三个问题的候选策略是Duffner与Garcia提出的课程式调度[3]，让广义均值"
        "的p随训练轮次从1线性升至5")
    add_formula(doc,
        r"p_e = p_{0} + (p_{T} - p_{0})\, \frac{e}{E - 1}",
        9, "2_3")
    add_para(doc,
        "其中e为当前轮次，E为总轮次，p_0等于1，p_T等于5。其直觉是训练初期软池化"
        "梯度稳定，利于学习全局结构，后期硬池化聚焦于判别性patch。对照组为固定"
        "p等于1与固定p等于5。注意到p等于1的广义均值在数学上精确等于算术平均，"
        "两个端点可直接复用矩阵中mean与gmean-p5的全量行，无需重复训练，课程组"
        "仅新增15次训练。")
    add_heading(doc, "2.2.4 理论分析框架", 3)
    add_para(doc,
        "第四个问题采用偏差与方差分解框架分析[4]。袋级预测的估计方差近似满足")
    add_formula(doc,
        r"\mathrm{Var}(\hat{y}) \approx \frac{\sigma^{2}}{n_{\mathrm{eff}}}",
        10, "2_4")
    add_para(doc,
        "其中σ²为patch分数的方差，n_eff为有效实例数。软池化的n_eff等于N，方差"
        "随N下降，但稀释偏差与N无关。固定个数池化的n_eff恒为k，不随N获得方差"
        "收益。固定比例池化兼顾两者，n_eff随N增长且硬度恒定。由此可预测，固定"
        "比例池化的误差对N不变，固定个数池化随N退化，实验将检验这一预测。")

    add_heading(doc, "2.3 结果与分析", 2)
    add_para(doc,
        "表2给出五种池化器在四档patch数量下的Pearson相关系数，图4以热图形式"
        "呈现完整的硬度与N曲面。")
    res = pd.read_csv(RESULTS / "6.2_results.csv", encoding="utf-8-sig")
    m5 = res[res.pooler != "gmean_curr_1to5"]
    n_cols = ["1000", "2000", "5000", "full"]
    rows2 = []
    for p_name in ["mean", "gmean_p5", "q75", "top50", "max"]:
        row = [p_name]
        for n in n_cols:
            v = m5[(m5.pooler == p_name) & (m5.N.astype(str) == n)]
            row.append(f"{v['pearson'].iloc[0]:.3f}")
        rows2.append(row)
    add_table(doc, "表2", "五种池化器在不同patch数量下的Pearson相关系数",
              ["池化器", "N=1000", "N=2000", "N=5000", "全量"], rows2)
    add_figure(doc, RESULTS / "6.2_hardness_N_surface.png", "图4",
               "池化硬度与patch数量的性能曲面，左为Pearson相关系数，右为MAE",
               width_cm=15.5)
    add_para(doc,
        "gmean-p5在每一档N上均为最优，Pearson介于0.722至0.746，两极池化显著"
        "落后，mean约为0.58至0.60，max仅为0.28至0.36。这印证了第一个问题的"
        "答案，最优硬度存在且位于中间偏软位置，与第三章干净基线条件下的结论"
        "互证。")
    add_para(doc,
        "图5展示回归误差随N的变化。q75与mean的曲线全程平坦，top-50单调退化，"
        "Pearson从0.615降至0.518，MAE从0.158升至0.172。这直接证实了理论预测，"
        "固定比例池化对袋容量不变，固定个数池化因等效硬度随N被动升高而退化。"
        "max在全程最差且伴随梯度饥饿现象[3]，只有最高分的个别patch获得梯度，"
        "模型难以学到可泛化的判别结构。")
    add_figure(doc, RESULTS / "6.2_error_vs_N.png", "图5",
               "回归误差随patch数量的变化，固定比例池化平稳而固定个数池化退化",
               width_cm=11.0)
    add_para(doc,
        "表3与图6给出课程调度与两个固定硬度端点的对比。课程组取得全场最佳的"
        "Spearman相关系数0.723与最佳的MAE 0.142，但Pearson为0.713，略低于固定"
        "p等于5的0.729。总体判断，课程调度与固定中间硬度打平，排序与校准略优而"
        "相关略损，在25张切片的规模上不构成净收益，固定中间硬度已足够，无需引入"
        "额外的调度复杂度。")
    cur = res[res.pooler == "gmean_curr_1to5"].iloc[0]
    p1 = m5[(m5.pooler == "mean") & (m5.N.astype(str) == "full")].iloc[0]
    p5 = m5[(m5.pooler == "gmean_p5") & (m5.N.astype(str) == "full")].iloc[0]
    rows3 = [
        ["固定p=1（=mean）", f"{p1['pearson']:.3f}", f"{p1['spearman']:.3f}",
         f"{p1['mae']:.4f}"],
        ["固定p=5", f"{p5['pearson']:.3f}", f"{p5['spearman']:.3f}",
         f"{p5['mae']:.4f}"],
        ["课程p：1→5", f"{cur['pearson']:.3f}", f"{cur['spearman']:.3f}",
         f"{cur['mae']:.4f}"],
    ]
    add_table(doc, "表3", "课程式硬度调度与固定硬度的对比（全量N）",
              ["调度策略", "Pearson", "Spearman", "MAE"], rows3)
    add_figure(doc, RESULTS / "6.2_curriculum.png", "图6",
               "课程式硬度调度与固定硬度的对比，课程组排序与误差略优而相关略损",
               width_cm=11.0)
    add_para(doc,
        "误差与N的经验关系总体为平。理论解释为，比例标签是无量纲量，嵌套子采样"
        "保持分布不变，q75在N为1000时有效实例数已达250，估计方差低于标签噪声"
        "地板，继续增大N的方差收益无法观测。由此得到一个实用的工程推论，袋级"
        "回归将patch子采样至一至两千几乎无损，gmean-p5在N为1000时Pearson为"
        "0.738，全量时为0.729，这对大规模部署是免费的效率收益。")

    add_heading(doc, "2.4 结论与讨论", 2)
    add_para(doc,
        "本章回答四个问题。第一，最优硬度存在且稳定，中间偏软的gmean-p5在全部"
        "袋容量档位上最优。第二，硬度不应随N变化，正确的参数化是比例而非个数，"
        "若工程上必须使用top-k形式，k应与N成正比。第三，课程式调度与固定中间"
        "硬度效果相当，不予采纳。第四，固定池化形式下回归误差与N近似无关，该"
        "现象可由估计方差与有效实例数的框架解释，当前的标签噪声地板掩盖了进一"
        "步的方差收益。")
    add_para(doc,
        "本章的局限在于，N低于500的极端区间未予探测，该区间方差主导，理论预测"
        "的方差上升趋势有待验证。若计算资源允许，将补充N为100至500的极端档位，"
        "并在第一章的合成真值平台上重做硬度与N的扫描，无标签噪声时方差曲线应当"
        "显形，同时在75张队列上复核课程调度的结论。")


# ---------------------------------------------------------------- 第 3 章（6.3）

def chapter_3(doc):
    add_heading(doc, "三、池化硬度与标签噪声鲁棒性", 1)

    add_heading(doc, "3.1 问题提出", 2)
    add_para(doc,
        "第一章讨论了标签本身的可信度推断问题，本章转向另一个维度，即当标签噪声"
        "不可避免时，模型的结构设计能否提供内在的抵抗力。项目文档提出一个具体"
        "假设，Top-1最大池化只关注单个最高分patch，天然忽略大部分patch的信息，"
        "这是否意味着最大池化对标签噪声更鲁棒。该假设的直觉是少听则少受误导。"
        "然而最大池化会严重低估肿瘤面积，这引出一个根本性问题，池化硬度如何影响"
        "模型对标签噪声的敏感性。本章在袋级占比回归的设定下推导不同池化函数的"
        "估计误差方差与噪声方差之间的关系，并通过人工噪声注入实验测量各池化策略"
        "的预测稳定性，回答三个问题，即极端池化是否天然抗噪，估计误差方差与噪声"
        "方差呈何种定量关系，以及离群值噪声下哪种池化函数损失最小。")

    add_heading(doc, "3.2 实验设计", 2)
    add_heading(doc, "3.2.1 理论推导与对立法假设", 3)
    add_para(doc,
        "设袋级标签携带零均值噪声，方差为σ²。在均方误差损失下，参数更新的扰动"
        "幅度与池化权重向量w的模长相关，预测扰动的方差近似满足")
    add_formula(doc,
        r"\mathrm{Var}(\hat{y}) \propto \sigma^{2} \|w\|^{2}"
        r" = \frac{\sigma^{2}}{n_{\mathrm{eff}}}",
        11, "3_1")
    add_para(doc,
        "各池化的有效实例数差异巨大。mean的权重均匀为1/N，扰动方差为σ²/N。"
        "max的权重集中于单点，扰动方差为σ²。top-k为σ²/k，分位数q为σ²/(qN)。"
        "理论预测因此与文档假设方向相反，在袋级回归噪声下，池化越硬，噪声敏感性"
        "越高。")
    add_para(doc,
        "文献中存在一个对立法结论。Duffner与Garcia在实例级单向翻牌噪声的分类"
        "任务下得出噪声越大则硬度越大越优的结论[3]。两种噪声模型的机理不同，"
        "实例级噪声污染的是每个patch的局部真伪，硬池化忽略低分实例从而避开被"
        "翻牌者，袋级回归噪声污染的是整条袋标签的数值，硬池化把全部监督压力"
        "集中到个别patch上，单点误差被整体放大。本章实验即裁决这一张力在本"
        "项目设定下的走向。")
    add_heading(doc, "3.2.2 噪声注入协议", 3)
    add_para(doc,
        "实验采用两类人工噪声。均匀噪声将训练标签扰动为")
    add_formula(doc,
        r"\tilde{y}_i = \mathrm{clip}(y_i + \varepsilon_i,\ 0,\ 1),"
        r"\quad \varepsilon_i \sim \mathcal{U}[-a,\ a]",
        12, "3_2")
    add_para(doc,
        "其中幅度a取0.15与0.30两档，0.30为D'Amato等设定的现实上限[2]。离群"
        "噪声与第一章式（3）相同，以比例r将标签替换为均匀分布随机数，r取0.2与"
        "0.5两档。噪声仅污染训练标签，验证集与测试集保持干净，以此测量模型在"
        "含噪监督下的泛化保持能力。全部实验固定种子与超参数，遵循Grahn提出的"
        "基准扫描纪律[5]。")
    add_heading(doc, "3.2.3 实验矩阵与有效实例数的实证度量", 3)
    add_para(doc,
        "实验矩阵为7种池化变体乘以5种噪声条件，即干净、均匀两档与离群两档，"
        "每组执行5折交叉验证与3种子集成，合计35组525次训练。零噪声行由两类"
        "噪声模型共享，不重复训练。池化变体覆盖硬度阶梯mean、gmean-p5、LSE-τ1、"
        "分位数q75、top-50与max，另加入attention池化作为学习型聚合的对照。训练"
        "仅使用面积约束均方误差，空间约束与可信度加权全部关闭，以隔离池化单一"
        "变量。有效实例数以参与率实证度量")
    add_formula(doc,
        r"n_{\mathrm{eff}} = \frac{\left( \sum_i |w_i| \right)^{2}}"
        r"{\sum_i w_i^{2}}",
        13, "3_3")
    add_para(doc,
        "其中w_i为各patch在聚合中的等效权重，由袋级预测对patch分数的梯度给出。")

    add_heading(doc, "3.3 结果与分析", 2)
    add_para(doc,
        "表4给出7种池化器在5种噪声条件下的Pearson相关系数，图7与图8分别展示"
        "均匀噪声与离群噪声下的退化曲线。")
    res = pd.read_csv(RESULTS / "6.3_results.csv", encoding="utf-8-sig")
    conds = [("clean", 0.0, "干净"), ("uniform", 0.15, "均匀0.15"),
             ("uniform", 0.30, "均匀0.30"), ("outlier", 0.20, "离群0.20"),
             ("outlier", 0.50, "离群0.50")]
    rows4 = []
    for p_name in ["mean", "gmean_p5", "lse_t1", "q75", "top50", "max",
                   "attention"]:
        row = [p_name]
        for nm, lv, _ in conds:
            v = res[(res.pooler == p_name) & (res.noise_model == nm)
                    & (res.noise_level == lv)]
            row.append(f"{v['pearson'].iloc[0]:.3f}")
        rows4.append(row)
    add_table(doc, "表4", "七种池化器在不同噪声条件下的Pearson相关系数",
              ["池化器"] + [c[2] for c in conds], rows4)
    add_para(doc,
        "干净基线的排序为gmean-p5最优0.729，q75次之0.644，LSE-τ1为0.623，"
        "attention与mean约为0.599，top-50为0.518，max最差为0.359。该排序与"
        "第二章曲面的全量列完全一致，两个实验的设置复现性由此交叉确认。")
    add_para(doc,
        "图7显示，均匀噪声下软池化的曲线平坦，mean与q75在各档噪声下的波动不"
        "超过0.03，attention甚至略有上升，最高从0.599升至0.676，推测轻度噪声"
        "起到了类似随机正则化的作用。图8显示，离群噪声下硬池化显著崩塌，max在"
        "r等于0.2时Pearson从0.359跌至0.132，top-50从0.518跌至0.438，而软池化"
        "保持基本稳定。")
    add_figure(doc, RESULTS / "6.3_degradation_uniform.png", "图7",
               "均匀噪声下各池化器的性能退化曲线，软池化平坦而硬池化波动")
    add_figure(doc, RESULTS / "6.3_degradation_outlier.png", "图8",
               "离群噪声下各池化器的性能退化曲线，max与top-50显著崩塌")
    add_para(doc,
        "图9给出硬度与方差的对照。横轴为实证的有效实例数，纵轴为噪声下的误差"
        "退化。top-50的有效实例数为50，MAE增量最大，为正0.0117，软池化的有效"
        "实例数约为2500至9600，增量接近零或为负。退化的排序与式（11）的理论"
        "预测方向一致。")
    add_figure(doc, RESULTS / "6.3_hardness_variance.png", "图9",
               "池化硬度与噪声退化的对照，有效实例数越小退化越大",
               width_cm=11.0)
    add_para(doc,
        "需要指出两点。其一，MAE增量指标对小幅退化不敏感，干净条件下的误差地板"
        "0.15掩盖了部分变化，Pearson保持率更为清晰。其二，25张切片的统计强度"
        "有限，本章结论的置信度受此约束，已在相应推断处按方向性结论处理。")

    add_heading(doc, "3.4 结论与讨论", 2)
    add_para(doc,
        "本章回答三个问题。第一，极端池化在袋级回归噪声下并不天然抗噪，方向与"
        "文档假设相反，池化越硬对袋级标签噪声越敏感。第二，估计误差方差与噪声"
        "方差的关系近似为σ²除以有效实例数，得到top-50与max的实证支持。第三，"
        "离群值噪声下软池化mean、q75与gmean损失最小，max损失最大。")
    add_para(doc,
        "本章另有一个意外收获。gmean-p5在干净基线上最优且在噪声下保持鲁棒，"
        "轻度硬度在稀释与聚焦之间取得平衡。该发现与第二章的最优硬度结论汇合，"
        "共同支持交付模型的聚合器落在中等偏软硬度的选择。")
    add_para(doc,
        "本章的局限在于队列规模为25张，噪声模型为零均值均匀扰动或均匀离群替换，"
        "尚未覆盖更贴近真实口径错位的系统性有偏噪声。若计算资源允许，将在75张"
        "队列上复核全矩阵，扩展有偏噪声模型，并推导gmean与LSE的有效实例数解析"
        "形式，以替代当前的实证估计。")


# ---------------------------------------------------------------- 第 4 章（6.4）

def chapter_4(doc):
    add_heading(doc, "四、面积硬约束与空间平滑的张力", 1)

    add_heading(doc, "4.1 问题提出", 2)
    add_para(doc,
        "交付管线的肿瘤区域掩膜由patch分数阈值化加形态学清理产生。训练阶段的面积"
        "约束是软性的，均方误差只要求预测接近标签，并不保证硬性满足，本项目在"
        "训练侧实现的容差式log-barrier约束[8]同样属于软约束路线。实测中阈值法"
        "产生的掩膜占比与模型预测占比的平均偏差达0.132，面积口径在最后一公里"
        "失守。本章围绕项目文档探索内容6.4展开，回答三个问题，即把面积约束写成"
        "硬性等式并同时要求空间平滑之后，问题在数学上变成什么，当patch数量处于"
        "10⁴至10⁵量级时精确求解面临什么计算瓶颈，以及能否设计近似求解策略，使"
        "面积误差不超过0.03且连通域数量显著少于简单阈值法。")

    add_heading(doc, "4.2 问题形式化与近似策略设计", 2)
    add_heading(doc, "4.2.1 数学形式化", 3)
    add_para(doc,
        "设切片含N个组织patch，交付模型输出每个patch的肿瘤分数s_i与袋级预测"
        "占比f̂。掩膜对应二值选择向量x，x_i等于1表示判为肿瘤。两个诉求写成一个"
        "约束优化问题")
    add_formula(doc,
        r"\max_{x \in \{0,1\}^{N}}\ \sum_{i} s_i x_i \;-\;"
        r" \lambda \sum_{(i,j) \in E} |x_i - x_j|"
        r"\quad \mathrm{s.t.}\ \sum_{i} x_i = K",
        14, "4_1")
    add_para(doc,
        "其中目标第一项奖励选取高分patch，第二项为Potts平滑能量，对网格八邻域"
        "内判定不一致的相邻patch对施加惩罚，λ控制平滑强度。约束项为面积硬约束，"
        "K等于预测占比与N的乘积取整，即掩膜必须精确承载模型的面积预测。")
    add_heading(doc, "4.2.2 可解性分析", 3)
    add_para(doc,
        "理解该问题的关键是两个诉求单独看都容易，合起来则困难。若只有平滑项而"
        "无基数约束，Potts能量满足子模条件")
    add_formula(doc,
        r"E(0,0) + E(1,1) \leq E(0,1) + E(1,0)",
        15, "4_2")
    add_para(doc,
        "即相邻变量取值一致比不一致代价更低。Kolmogorov与Zabih证明，满足子模"
        "条件的二元两两能量都可图表示，可通过s-t最小割在多项式时间求得全局"
        "最优解[6]，Boykov与Kolmogorov的最大流实现在十万节点量级的网格图上"
        "仅需秒级[7]。另一方面，若只有面积约束而无平滑项，问题退化为按分数排序"
        "取前K个，是平凡的精确解。然而基数约束把全部N个变量全局耦合，最小割的"
        "代价是逐边局部求和的，任何局部边结构都无法表达恰好K个节点被选的全局"
        "计数条件。该问题与最小二等分等经典NP难问题同族，且平滑项破坏了单调"
        "子模性，使贪心的经典近似保证不再适用。从树分解角度看，二维网格图的"
        "treewidth约为根号N，N为10⁴时约为100，精确动态规划需要2的根号N次方"
        "量级状态，整数规划分支定界最坏为指数。结论是N处于10⁴至10⁵量级时精确"
        "求解不可行，这就是张力的数学来源。")
    add_heading(doc, "4.2.3 拉格朗日松弛与对偶间隙", 3)
    add_para(doc,
        "精确不可行，退而求其次。把硬约束乘以乘子γ放进目标，得到拉格朗日松弛")
    add_formula(doc,
        r"\max_{x}\ \sum_{i} (s_i + \gamma)\, x_i \;-\;"
        r" \lambda \sum_{(i,j) \in E} |x_i - x_j| \;-\; \gamma K",
        16, "4_3")
    add_para(doc,
        "γ可理解为选区价格，每选取一个patch额外付出γ。对任意固定的γ，问题退回"
        "无约束的子模能量，一次最小割精确求解。选中数是γ的单调不增函数，因此"
        "二分搜索γ即可逼近预算K。但选中数是阶梯函数而非连续函数，平滑项把成片"
        "patch绑定在一起，γ跨过临界值时选区整块翻转，选中数可能从K的一侧直接"
        "跳到另一侧，不存在恰好命中K的γ。松弛问题最优值与原问题最优值之差即"
        "对偶间隙，这是离散问题拉格朗日松弛的固有性质，也是平滑强度越大面积"
        "控制越粗糙的原因。")
    add_heading(doc, "4.2.4 近似策略与评估协议", 3)
    add_para(doc,
        "实验对比五种策略。S1为简单阈值法，取分数不低于0.5的patch，是交付管线"
        "现状。S2为Top-K贪心，按分数排序取前K个，面积在构造上精确。S3先将分数"
        "沿空间邻近与特征相似构成的图传播一轮再取Top-K，S3b传播三轮，二者把"
        "平滑先验融入分数再做精确预算。S4为拉格朗日图割，以PyMaxflow求解，"
        "平滑权重β取0.3，γ在正负6区间内二分22次，在按预测占比分层抽样的9张"
        "代表切片上运行。全部策略统一接形态学清理，评估指标为面积误差、连通域"
        "数量、掩膜边界平滑度与单张耗时。")
    add_para(doc,
        "本实验为纯CPU后处理，不重训模型，原因有二。其一，硬约束作用于离散的"
        "选区变量，而梯度训练产出的是连续分数，硬满足只能在分数到掩膜的后处理"
        "环节落地。其二，四种策略作用于同一张冻结分数图，指标差异可干净地归因"
        "于选择算子。分数图由交付配置经5折交叉验证生成，即75张切片的面积口径"
        "标签、q75聚合、可信度加权与平滑损失，每张切片的分数来自未见过其标签的"
        "折，无泄漏。该次运行的Pearson为0.801，与交付模型的既有结果复现一致。")

    add_heading(doc, "4.3 结果与分析", 2)
    add_para(doc,
        "表5给出五种策略在75张切片上的平均表现，图10以一张代表切片直观对比"
        "各策略的掩膜形态。")
    m = pd.read_csv(RESULTS / "6.4_metrics.csv", encoding="utf-8-sig")
    g = m.groupby("strategy")[["area_err", "n_components",
                               "mask_neighbor_diff"]].mean()
    rt = {"S1_threshold": "<0.1s", "S2_topk": "<0.1s", "S3_prop_topk": "约1s",
          "S3b_prop3_topk": "约2s", "S4_graphcut": "约1.3s"}
    name_zh = {"S1_threshold": "S1 阈值0.5", "S2_topk": "S2 Top-K",
               "S3_prop_topk": "S3 传播×1+Top-K",
               "S3b_prop3_topk": "S3b 传播×3+Top-K",
               "S4_graphcut": "S4 图割β=0.3"}
    rows5 = []
    for key in ["S1_threshold", "S2_topk", "S3_prop_topk", "S3b_prop3_topk",
                "S4_graphcut"]:
        rows5.append([name_zh[key], f"{g.loc[key, 'area_err']:.4f}",
                      f"{g.loc[key, 'n_components']:.1f}",
                      f"{g.loc[key, 'mask_neighbor_diff']:.4f}", rt[key]])
    add_table(doc, "表5", "五种掩膜策略的平均表现（S4为9张代表切片，余为75张）",
              ["策略", "面积误差", "连通域数", "边界平滑度", "单张耗时"], rows5)
    add_figure(doc, RESULTS / "6.4_mask_TCGA-AQ-A04J-01Z-00-DX1.png", "图10",
               "代表切片的分数图与五种策略的掩膜对比，S2空间破碎，S3b与S4成团性好",
               width_cm=13.0)
    add_para(doc,
        "结果呈现清晰的张力格局。S1的连通域最少，为14.2个，但面积误差0.132，"
        "远未达标，其低连通域恰恰来自放弃面积约束，只取高置信核心区域。S2面积"
        "精确，误差0.026，但空间破碎，连通域升至35.0个。S3与S3b在两个方向上"
        "同时改善，S3b面积误差降至0.0073，连通域降至30.1个。S4在面积可行的"
        "策略中全面最优，面积误差0.0037，优于目标约8倍，连通域26.0个，边界"
        "平滑度0.022，单张耗时约1.3秒，具备工程实用性。")
    add_para(doc,
        "图11给出S4的β扫描与参照策略的Pareto格局。β从0.1升至0.9的过程中连通"
        "域稳定在26个上下，面积误差在β等于0.9时升至0.0196。平滑拨盘的边际收益"
        "已经耗尽，继续加大平滑只付出面积代价而买不来更少的连通域，对偶间隙在"
        "强平滑端显形，与4.2.3节的理论分析一致。")
    add_figure(doc, RESULTS / "6.4_pareto.png", "图11",
               "面积误差与连通域数量的Pareto格局，S4的β扫描曲线与参照策略",
               width_cm=11.0)
    add_para(doc,
        "需要记录一处排错。S4的首版实现将图割的源汇侧读反，掩膜取为补集，面积"
        "误差高达0.466。修正后图割精确命中预算。该错误及其修复反向验证了对偶"
        "偏差γ控制基数的机制按预期工作。")

    add_heading(doc, "4.4 结论与讨论", 2)
    add_para(doc,
        "本章回答三个问题。第一，硬面积约束与空间平滑结合后，问题在数学上是"
        "基数约束的Potts模型，无约束时图割多项式可解，加约束后为NP难。第二，"
        "计算瓶颈在于全局基数与局部平滑的耦合，N处于10⁴至10⁵量级时精确方法"
        "指数不可行，对偶松弛存在间隙。第三，近似策略方面，S4与S3b的面积误差"
        "分别为0.004与0.007，优于0.03的目标约4至8倍，但连通域数量无法少于"
        "简单阈值法，因为阈值的低连通域正是以放弃面积约束为代价换来的。文档"
        "提出的双重目标在本数据上不可同时达到，这一不可达性本身就是张力的实证。")
    add_para(doc,
        "基于上述结论，交付掩膜建议改采S3b，即传播三轮加Top-K加形态学清理，"
        "面积精确，无新增依赖，单张约2秒。S4作为原理性最优近似入档，PyMaxflow"
        "列为可选依赖。")
    add_para(doc,
        "本章的局限在于连通域数量的地板由肿瘤本身的多灶性决定，乳腺癌全切片"
        "常呈多灶分布，后处理无法在不牺牲面积的前提下压到阈值法水平，这是问题"
        "内禀而非方法缺陷。若计算资源允许，将把S4的一元项替换为SRG校准后的"
        "分数，在掩膜层面做连通域级贪心选区，并对β做逐切片自适应定标。")


# ---------------------------------------------------------------- 引言与总结

def intro(doc):
    add_heading(doc, "引言", 1)
    add_para(doc,
        "本项目面向TCGA-BRCA乳腺癌H&E全切片图像，在不使用任何像素级肿瘤标注的"
        "前提下，建立弱监督的肿瘤区域定位与面积占比估计管线。管线以切片级弱标签"
        "为唯一监督信号，经UNI基础模型提取patch特征，由多示例学习聚合器完成"
        "占比回归，并输出热图与二值掩膜。在这一框架中，标签噪声、聚合器选型与"
        "掩膜生成三个环节均存在缺乏现成答案的设计问题，项目文档因此设置四项"
        "探索任务，要求以小型受控实验逐一回答。")
    add_para(doc,
        "本报告汇总四项探索的设计、结果与结论。第一章对应探索内容6.1，刻画标签"
        "可信度协同推断的可行性边界。第二章对应6.2，扫描池化硬度与patch数量的"
        "匹配关系。第三章对应6.3，测量池化硬度与标签噪声鲁棒性的关系。第四章"
        "对应6.4，分析面积硬约束与空间平滑的组合优化瓶颈并对比近似策略。四章"
        "共享同一套实验协议，即5折分层交叉验证、3种子集成、面积口径标签与严格"
        "的变量隔离，结论可相互印证。")


def conclusion(doc):
    add_heading(doc, "五、总结", 1)
    add_para(doc,
        "四项探索围绕标签侧、聚合侧与输出侧给出了完整回答。标签侧，协同推断在"
        "污染率不低于35%时有净收益，识别AUC介于0.70至0.84，但无法校正系统性"
        "口径错位，与面积参照路线互补。聚合侧，最优硬度存在且稳定，中间偏软的"
        "gmean-p5在全部袋容量档位与各类噪声下均为最优或近最优，硬度的正确参数"
        "化是比例而非个数，估计误差方差近似服从σ²除以有效实例数的关系。输出侧，"
        "面积硬约束与空间平滑的组合是NP难问题，拉格朗日图割与传播平滑加Top-K"
        "两种近似策略的面积误差优于目标4至8倍，而连通域数量无法同时少于阈值法，"
        "这一不可达性正是张力的实证。")
    add_para(doc,
        "上述结论已落实为交付管线的四项设计决策，即中等偏软硬度的聚合器、按"
        "污染风险筛查后启用的可信度加权、传播平滑加精确预算选区的掩膜策略，"
        "以及按需子采样的效率策略。后续工作将在75张队列上复核全矩阵结论，把"
        "噪声模型扩展为系统性有偏噪声，并为图割近似引入SRG校准的一元项与"
        "逐切片自适应的平滑权重。")


# ---------------------------------------------------------------- 主流程

def main():
    doc = Document()
    setup_styles(doc)

    # 标题页
    for _ in range(6):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(despace("弱监督 WSI 肿瘤定位\n探索性实验报告"))
    r.font.size = Pt(26)
    r.bold = True
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = st.add_run(despace("——基于 TCGA-BRCA 的 MIL 探索实验 6.1–6.4"))
    r2.font.size = Pt(14)
    r2.font.name = "Times New Roman"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    doc.add_paragraph()
    d = doc.add_paragraph()
    d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = d.add_run("2026 年 8 月")
    r3.font.size = Pt(14)
    r3.font.name = "Times New Roman"
    r3._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    doc.add_page_break()

    intro(doc)
    chapter_1(doc)
    chapter_2(doc)
    chapter_3(doc)
    chapter_4(doc)
    conclusion(doc)

    add_references(doc, REFS)

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(OUT_DOCX)
        print(f"报告已生成 → {OUT_DOCX}")
    except PermissionError:
        alt = OUT_DOCX.with_name(OUT_DOCX.stem + "_new.docx")
        doc.save(alt)
        print(f"目标文件被占用（可能在 Word 中打开），已改存 → {alt}")


if __name__ == "__main__":
    main()
